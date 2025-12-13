"""
Extrator de texto de documentos PDF, DOCX e TXT
"""
import re
import chardet
from pathlib import Path
from typing import Optional
import structlog

logger = structlog.get_logger()


class TextExtractor:
    """Extrai e limpa texto de diferentes formatos de documento"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def extract(self, file_path: str) -> str:
        """
        Extrai texto do arquivo baseado na extensão
        """
        path = Path(file_path)
        extension = path.suffix.lower().replace('.', '')
        
        if extension not in self.supported_formats:
            raise ValueError(f"Formato não suportado: {extension}")
        
        logger.info("text_extraction_started", file=file_path, format=extension)
        
        extractors = {
            'pdf': self._extract_pdf,
            'docx': self._extract_docx,
            'txt': self._extract_txt
        }
        
        raw_text = extractors[extension](file_path)
        cleaned_text = self._clean_text(raw_text)
        
        logger.info(
            "text_extraction_completed",
            file=file_path,
            raw_length=len(raw_text),
            cleaned_length=len(cleaned_text)
        )
        
        return cleaned_text
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extrai texto de PDF usando pdfplumber (melhor para layouts complexos)"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return '\n\n'.join(text_parts)
        
        except Exception as e:
            logger.warning("pdfplumber_failed", error=str(e), fallback="PyPDF2")
            return self._extract_pdf_fallback(file_path)
    
    def _extract_pdf_fallback(self, file_path: str) -> str:
        """Fallback usando PyPDF2"""
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    def _extract_docx(self, file_path: str) -> str:
        """Extrai texto de DOCX preservando estrutura"""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Também extrai de tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return '\n\n'.join(text_parts)
    
    def _extract_txt(self, file_path: str) -> str:
        """Extrai texto de TXT com detecção de encoding"""
        # Detecta encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
        
        # Lê com encoding detectado
        with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
            return f.read()
    
    def _clean_text(self, text: str) -> str:
        """
        Limpa o texto removendo artefatos comuns
        """
        if not text:
            return ""
        
        # Remove múltiplas quebras de linha
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove espaços múltiplos
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove números de página comuns
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        text = re.sub(r'^\s*Página\s+\d+\s*$', '', text, flags=re.MULTILINE | re.IGNORECASE)
        text = re.sub(r'^\s*Page\s+\d+\s*$', '', text, flags=re.MULTILINE | re.IGNORECASE)
        
        # Remove headers/footers repetitivos (linhas curtas repetidas)
        lines = text.split('\n')
        line_counts = {}
        for line in lines:
            stripped = line.strip()
            if len(stripped) < 50:  # Linhas curtas podem ser headers/footers
                line_counts[stripped] = line_counts.get(stripped, 0) + 1
        
        # Remove linhas que aparecem mais de 3 vezes
        repeated_lines = {line for line, count in line_counts.items() if count > 3}
        lines = [line for line in lines if line.strip() not in repeated_lines]
        text = '\n'.join(lines)
        
        # Remove caracteres de controle
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        # Normaliza espaços em branco
        text = text.strip()
        
        return text


class ContentValidator:
    """Valida se o conteúdo é suficiente para geração de questões"""
    
    def __init__(self, min_words: int = 500):
        self.min_words = min_words
    
    def validate(self, text: str) -> dict:
        """
        Valida o conteúdo e retorna análise
        """
        words = text.split()
        word_count = len(words)
        
        # Detecção simples de idioma (verifica palavras comuns em português)
        pt_words = {'de', 'da', 'do', 'que', 'e', 'em', 'um', 'uma', 'para', 'com', 'não', 'os', 'as'}
        text_words_lower = set(w.lower() for w in words[:200])  # Amostra inicial
        pt_matches = len(pt_words.intersection(text_words_lower))
        language = 'pt-BR' if pt_matches >= 5 else 'unknown'
        
        is_sufficient = word_count >= self.min_words
        
        suggestions = []
        if not is_sufficient:
            suggestions.append(
                f"O texto possui {word_count} palavras. "
                f"Recomendamos pelo menos {self.min_words} palavras para geração de questões de qualidade."
            )
        
        if language == 'unknown':
            suggestions.append(
                "O idioma do texto não foi identificado claramente como português. "
                "Isso pode afetar a qualidade das questões geradas."
            )
        
        return {
            'word_count': word_count,
            'language': language,
            'is_sufficient': is_sufficient,
            'suggestions': suggestions
        }
